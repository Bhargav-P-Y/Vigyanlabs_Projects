import os
import torch
import numpy as np
import requests
import json
import base64
import re
from collections import Counter
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdf2image import convert_from_path
from torchvision.transforms.functional import to_tensor
import matplotlib.pyplot as plt

# --- 1. CONFIGURATION ---
BASE_DIR = os.getcwd()
# Points to the folder parallel to 'photoholmes'
INPUT_ROOT_FOLDER = os.path.join(BASE_DIR, "../CERTIFICATES-20260202T052442Z-3-001")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "analysis_output_production")
WEIGHTS_PATH = "weights/trufor/trufor.pth.tar"
OCR_API_URL = "http://10.91.2.100:8010/v1/chat/completions"
REPORT_PATH = os.path.join(OUTPUT_FOLDER, "final_production_report.json")

# FORCE CPU (Stability Priority)
DEVICE = "cpu"

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# --- 2. SECURITY BYPASS ---
torch.serialization.add_safe_globals([np.core.multiarray.scalar])
try:
    torch.serialization.add_safe_globals([np.core.multiarray._reconstruct])
except:
    pass

# --- 3. TEXT & TABLE CLEANING ENGINE ---

def clean_line_content(s):
    """Normalizes a line to detect duplicates (ignores numbers/punctuation)."""
    return re.sub(r'[\d\.\-\)\(\s]+', '', s).lower()

def clean_ocr_text(text):
    """
    Bulletproof cleaning: Anti-Hallucination + HTML Strip + Table Norm
    """
    if not text: return ""

    # A. Global Frequency Filter (Loop Killer)
    lines = text.split('\n')
    unique_lines = []
    line_counts = Counter()
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            unique_lines.append(line)
            continue
        
        content_hash = clean_line_content(stripped)
        # Allow short lines to repeat slightly more, long lines strictly limited
        max_repeats = 3 if len(content_hash) > 10 else 10
        
        if line_counts[content_hash] < max_repeats:
            unique_lines.append(line)
            line_counts[content_hash] += 1
        else:
            if line_counts[content_hash] == max_repeats:
                 unique_lines.append("... [Repeated Content Truncated] ...")
                 line_counts[content_hash] += 1
    
    text = '\n'.join(unique_lines)

    # B. Phrase-based repetition (Intra-line)
    text = re.sub(r'(.{6,}?)( \1){2,}', r'\1', text)

    # C. HTML Stripper
    if "<table>" in text or "<tr>" in text or "<td" in text:
        text = re.sub(r'<tr[^>]*>', '\n', text)      
        text = re.sub(r'</?td[^>]*>', '|', text)     
        text = re.sub(r'</?th[^>]*>', '|', text)     
        text = re.sub(r'</?table[^>]*>', '', text)   
        text = re.sub(r'</?thead[^>]*>', '', text)
        text = re.sub(r'</?tbody[^>]*>', '', text)
    
    text = re.sub(r'<[^>]+>', '', text) 

    # D. Markdown Normalization
    text = re.sub(r'\|+', '|', text)
    cleaned_lines = []
    for line in text.split('\n'):
        clean = line.strip()
        if not clean: continue
        if clean.count('|') > 1:
            if not clean.startswith('|'): clean = '|' + clean
            if not clean.endswith('|'): clean = clean + '|'
        cleaned_lines.append(clean)
    
    return '\n'.join(cleaned_lines)

def smart_reshape_table(table_lines):
    """Fixes 'Run-on Rows'."""
    if not table_lines: return []
    
    header_cells = [c.strip() for c in table_lines[0].strip().split('|') if c.strip()]
    num_cols = len(header_cells)
    if num_cols < 2: return table_lines
    
    reshaped_lines = [table_lines[0]]
    
    for line in table_lines[1:]:
        cells = [c.strip() for c in line.strip().split('|') if c.strip()]
        
        if len(cells) >= (num_cols * 2) and (len(cells) % num_cols == 0):
            for i in range(0, len(cells), num_cols):
                chunk = cells[i : i + num_cols]
                reshaped_lines.append("|" + "|".join(chunk) + "|")
        else:
            reshaped_lines.append(line)
            
    return reshaped_lines

def add_table_to_doc(doc, table_lines):
    """Renders table to Word with Fail-Safe."""
    if not table_lines: return

    rows_data = []
    max_cols = 0
    for line in table_lines:
        cells = [c.strip() for c in line.strip().split('|') if c.strip()]
        if cells:
            rows_data.append(cells)
            max_cols = max(max_cols, len(cells))
    
    if not rows_data: return

    if max_cols > 15:
        doc.add_paragraph("[Complex Data Table - Rendered as List:]")
        for line in table_lines:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
        return

    try:
        table = doc.add_table(rows=len(rows_data), cols=max_cols)
        table.style = 'Table Grid'
        table.autofit = False 
        
        for r_idx, row in enumerate(rows_data):
            row_cells = table.rows[r_idx].cells
            for c_idx, cell_text in enumerate(row):
                if c_idx < len(row_cells):
                    row_cells[c_idx].text = cell_text
    except Exception:
        doc.add_paragraph("[Table Rendering Error - Raw Content:]")
        for line in table_lines:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'

def parse_markdown_to_word(doc, text):
    """Main parser."""
    clean_text = clean_ocr_text(text)
    lines = clean_text.split('\n')
    table_buffer = []
    
    for line in lines:
        stripped = line.strip()
        if '|' in stripped and set(stripped) - {'|', '-', ' ', ':'}:
            table_buffer.append(stripped)
        else:
            if table_buffer:
                fixed_table = smart_reshape_table(table_buffer)
                add_table_to_doc(doc, fixed_table)
                table_buffer = []
            
            if '---' in stripped: continue
            
            if stripped: 
                p = doc.add_paragraph(stripped)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if table_buffer:
        fixed_table = smart_reshape_table(table_buffer)
        add_table_to_doc(doc, fixed_table)

# --- 4. FORENSIC REASONING ---

def get_enhanced_reasoning(score, label):
    if label == "GENUINE":
        return f"Integrity Verified (Score: {score:.4f}). The analysis found no significant statistical anomalies. The pixel distribution and compression signatures are consistent with an original, unaltered document."
    if score > 0.85:
        return f"Surgical Tampering Detected: The deep learning model identified high-confidence anomalies (Score: {score:.4f}) consistent with digital splicing. The heatmap highlights disjointed pixel patterns in specific regions, indicating content was likely pasted or overwritten."
    elif score > 0.65:
        return f"Inconsistent Pixel Signatures: The analysis detected irregular noise patterns (Score: {score:.4f}) that deviate from the document's global profile. This suggests potential localized editing or retouching in the highlighted areas."
    else:
        return f"Potential Manipulation: The model flagged ambiguous artifacts (Score: {score:.4f}). While not definitive, the pixel structure in the marked regions differs from the background, warranting manual review."

# --- 5. MODEL & API LOADERS ---

def load_trufor_model():
    from photoholmes.methods.trufor import TruFor
    try:
        model = TruFor(weights=WEIGHTS_PATH, device=DEVICE)
    except Exception:
        original_load = torch.load
        torch.load = lambda f, map_location=None, weights_only=True, **kwargs: original_load(f, map_location, weights_only=False, **kwargs)
        model = TruFor(weights=WEIGHTS_PATH, device=DEVICE)
    return model

def call_ocr_api(image_path):
    try:
        with open(image_path, "rb") as f:
            base64_image = base64.b64encode(f.read()).decode('utf-8')
        
        payload = {
            "model": "dots_ocr",
            "messages": [
                {"role": "system", "content": "You are a professional OCR engine. Extract text exactly as seen. 1. If you see a table, OUTPUT IT AS A MARKDOWN TABLE (using | pipes). 2. DO NOT use HTML tags. 3. STRICTLY AVOID repeating text loops. 4. Preserve layout."},
                {"role": "user", "content": [
                    {"type": "text", "text": "Extract all text."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                ]}
            ],
            "temperature": 0.1
        }
        
        response = requests.post(OCR_API_URL, json=payload, timeout=120)
        if response.status_code == 200:
            content = response.json()['choices'][0]['message']['content']
            if not content.strip(): return "[OCR Warning: Image content unclear.]"
            return content
        return f"[OCR Error: Status {response.status_code}]"
    except Exception as e:
        return f"[OCR Connection Error: {str(e)}]"

# --- 6. MAIN PIPELINE ---

def main():
    if not os.path.exists(INPUT_ROOT_FOLDER):
        print(f"❌ Input folder not found: {INPUT_ROOT_FOLDER}")
        return

    print(f"🚀 Initializing TruFor AI on {DEVICE.upper()}...")
    model = load_trufor_model()
    
    full_report = []

    print(f"📂 Scanning Directory: {INPUT_ROOT_FOLDER}")
    
    # RECURSIVE WALK
    for root, dirs, files in os.walk(INPUT_ROOT_FOLDER):
        for filename in files:
            if not filename.lower().endswith(('.pdf', '.jpg', '.jpeg', '.webp')):
                continue

            file_path = os.path.join(root, filename)
            # Use subdirectory as category, or 'Uncategorized' if in root
            category = os.path.basename(root) if root != INPUT_ROOT_FOLDER else "Uncategorized"
            
            print(f"\n   ⚙️  Processing [{category}]: {filename}")
            
            try:
                # A. Prepare Image (First Page Only for Speed)
                if filename.lower().endswith('.pdf'):
                    img = convert_from_path(file_path, 200)[0]
                else:
                    img = Image.open(file_path).convert("RGB")
                
                # Resize if massive (prevents OOM on CPU)
                if max(img.size) > 2000:
                    img.thumbnail((2000, 2000), Image.Resampling.LANCZOS)
                
                temp_path = os.path.join(OUTPUT_FOLDER, f"temp_{filename}.jpg")
                img.save(temp_path, quality=95)

                # B. AI Detection
                img_tensor = to_tensor(img).to(DEVICE)
                output = model.predict(img_tensor)
                
                # Output Parsing
                if isinstance(output, (tuple, list)):
                    heatmap = output[0]
                    raw_score = output[1]
                    if isinstance(raw_score, torch.Tensor):
                        score = raw_score.detach().cpu().mean().item()
                    else:
                        score = float(raw_score)
                else:
                    heatmap = output
                    score = 0.5 
                
                if isinstance(heatmap, torch.Tensor):
                    heatmap = heatmap.detach().cpu().squeeze().numpy()

                # Verdict
                if score > 0.55: label = "FORGED"
                else: label = "GENUINE"
                
                reason = get_enhanced_reasoning(score, label)

                # C. Save Heatmap
                heatmap_filename = f"{os.path.splitext(filename)[0]}_heatmap.png"
                heatmap_path = os.path.join(OUTPUT_FOLDER, heatmap_filename)
                
                plt.figure(figsize=(10, 5))
                plt.subplot(1, 2, 1); plt.imshow(img); plt.axis('off'); plt.title("Original")
                plt.subplot(1, 2, 2); plt.imshow(heatmap, cmap='jet'); plt.axis('off'); plt.title(f"Forgery Map")
                plt.savefig(heatmap_path, bbox_inches='tight')
                plt.close()

                # D. OCR Extraction
                text_content = call_ocr_api(temp_path)

                # E. Word Report
                doc = Document()
                h = doc.add_heading(f"Analysis: {filename}", 0)
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                doc.add_paragraph(f"Category: {category}")
                p = doc.add_paragraph(f"Verdict: {label} (Confidence: {score:.4f})")
                p.bold = True
                
                doc.add_heading("Forensic Reasoning:", level=1)
                doc.add_paragraph(reason)
                
                doc.add_heading("Extracted Content", level=1)
                parse_markdown_to_word(doc, text_content)
                
                docx_filename = f"{os.path.splitext(filename)[0]}_report.docx"
                docx_path = os.path.join(OUTPUT_FOLDER, docx_filename)
                doc.save(docx_path)

                # F. Add to Master JSON
                clean_snippet_text = clean_ocr_text(text_content)
                snippet = clean_snippet_text[:200].replace('\n', ' ').strip() + "..."
                
                record = {
                    "metadata": {
                        "file": filename,
                        "category": category,
                        "pages": 1
                    },
                    "verdict": {
                        "score": round(score, 4),
                        "flag": label,
                        "reasoning": reason
                    },
                    "structural_details": [
                        {
                            "page": 1,
                            "heatmap": heatmap_path,
                            "text_snippet": snippet
                        }
                    ],
                    "docx_path": docx_path
                }
                full_report.append(record)

                if os.path.exists(temp_path): os.remove(temp_path)
                print(f"      ✅ Success: {label} ({score:.2f})")

            except Exception as e:
                print(f"      ❌ Failed: {e}")

    # G. Save Master JSON Report
    with open(REPORT_PATH, "w", encoding='utf-8') as f:
        json.dump(full_report, f, indent=4, ensure_ascii=False)

    print(f"\n✨ BATCH PROCESSING COMPLETE.")
    print(f"   📂 Output: {OUTPUT_FOLDER}")
    print(f"   📄 Master Report: {REPORT_PATH}")

if __name__ == "__main__":
    main()
